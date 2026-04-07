#!/usr/bin/env python3
"""
import_fdr.py — Import des anciennes FDR vers Supabase
Formats : .docx, .doc, .odt, .pdf

Usage :
    python3 import_fdr.py /chemin/vers/dossier
    python3 import_fdr.py /chemin/vers/dossier --prod   (insertion réelle)
"""

import os, re, sys, json, zipfile, requests
from pathlib import Path

SUPABASE_URL = "https://whlxbfnmyqdflmxosfse.supabase.co"
SUPABASE_KEY = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJpc3MiOiJzdXBhYmFzZSIsInJlZiI6IndobHhiZm5teXFkZmxteG9zZnNlIiwicm9sZSI6ImFub24iLCJpYXQiOjE3NzI3ODA5MTksImV4cCI6MjA4ODM1NjkxOX0.vf3sdnJRnnXyIx998fhPSIUPX0WS7KqDbvAwesCzOcE"
MODE_PROD = "--prod" in sys.argv

# ══ EXTRACTEURS ══════════════════════════════════════════════

def extraire_odt(p):
    try:
        from odf import teletype
        from odf.opendocument import load
        return teletype.extractText(load(p).text)
    except Exception as e:
        return ""

def extraire_docx(p):
    try:
        from docx import Document
        doc = Document(p)
        lignes = []
        for para in doc.paragraphs:
            if para.text.strip():
                lignes.append(para.text.strip())
        # Tables aussi
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lignes.append(cell.text.strip())
        return "\n".join(lignes)
    except Exception as e:
        return ""

def extraire_doc(p):
    """Tente d'extraire un .doc via antiword ou catdoc"""
    try:
        import subprocess
        for cmd in [['antiword', str(p)], ['catdoc', str(p)]]:
            try:
                r = subprocess.run(cmd, capture_output=True, text=True, timeout=10)
                if r.stdout.strip():
                    return r.stdout
            except FileNotFoundError:
                continue
        # Fallback: lire les chaînes UTF-8 du binaire
        with open(p, 'rb') as f:
            data = f.read()
        texte = data.decode('latin-1', errors='ignore')
        lignes = [l.strip() for l in texte.split('\n') if len(l.strip()) > 4 and l.strip().isprintable()]
        return "\n".join(lignes[:100])
    except Exception as e:
        return ""

def extraire_pdf(p):
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(p) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        return "\n".join(pages)
    except Exception as e:
        return ""

# ══ UTILITAIRES ═══════════════════════════════════════════════

MOIS_FR = {
    'janvier':'01','fevrier':'02','février':'02','mars':'03','avril':'04',
    'mai':'05','juin':'06','juillet':'07','aout':'08','août':'08',
    'septembre':'09','octobre':'10','novembre':'11','decembre':'12','décembre':'12'
}

def normaliser_date(s):
    """Convertit diverses formes de date en YYYY-MM-DD"""
    s = s.strip()
    # DD/MM/YYYY ou DD-MM-YYYY
    m = re.match(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{4})', s)
    if m: return f"{m.group(3)}-{m.group(2).zfill(2)}-{m.group(1).zfill(2)}"
    # YYYY-MM-DD
    m = re.match(r'(\d{4})[/\-](\d{2})[/\-](\d{2})', s)
    if m: return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    # DD mois YYYY
    m = re.match(r'(\d{1,2})\s+(\w+)\s+(\d{4})', s, re.IGNORECASE)
    if m:
        mois = MOIS_FR.get(m.group(2).lower())
        if mois: return f"{m.group(3)}-{mois}-{m.group(1).zfill(2)}"
    return None

def date_du_nom(nom):
    """Extrait une date depuis le nom de fichier"""
    # YYYY-MM-DD en début
    m = re.match(r'(\d{4})-(\d{2})-(\d{2})', nom)
    if m: return f"{m.group(1)}-{m.group(2)}-{m.group(3)}"
    # DD_mois_YYYY
    m = re.search(r'(\d{1,2})[_\s](\w+)[_\s](\d{4})', nom)
    if m:
        mois = MOIS_FR.get(m.group(2).lower())
        if mois: return f"{m.group(3)}-{mois}-{m.group(1).zfill(2)}"
    return None

def chercher(texte, patterns, cast=None):
    """Cherche un pattern et retourne la valeur nettoyée"""
    for pat in patterns:
        m = re.search(pat, texte, re.IGNORECASE | re.MULTILINE)
        if m:
            val = m.group(1).strip().rstrip('.')
            val = re.sub(r'\s+', ' ', val).strip()
            if not val: continue
            if cast == float:
                try:
                    n = re.search(r'[\d,\.]+', val.replace(' ',''))
                    return float(n.group().replace(',','.')) if n else None
                except: continue
            if cast == int:
                try:
                    n = re.search(r'\d+', val.replace(' ','').replace('\u202f',''))
                    return int(n.group()) if n else None
                except: continue
            return val
    return None

def normaliser_heure(s):
    if not s: return None
    m = re.match(r'(\d{1,2})h(\d{0,2})', s.strip(), re.IGNORECASE)
    if m:
        hh = m.group(1).zfill(2)
        mm = (m.group(2) or '00').zfill(2)
        return f"{hh}:{mm}"
    return None

# ══ PARSER ════════════════════════════════════════════════════

def parser(texte, nom_fichier):
    f = {k: None for k in [
        'date_rando','nom_rando','animateur','parking_covoit','heure_rv',
        'parking_depart','gps','distance','denivele','duree',
        'ibp','effort','technicite','risque','couts','remarques','profil_png'
    ]}

    # ── DATE ──
    # 1. Depuis le nom de fichier (le plus fiable)
    f['date_rando'] = date_du_nom(nom_fichier)
    if not f['date_rando']:
        # 2. Depuis le texte : patterns variés
        for pat in [
            r'^((?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)\s+\d{1,2}[/\-]\d{1,2}[/\-]\d{4})',
            r'^(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})',
            r'^((?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)\s+\d{1,2}\s+\w+\s+\d{4})',
            r'^(\d{1,2}\s+\w+\s+\d{4})',
            r'[Dd]ate\s*:?\s*(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})',
            r'(?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)\s+(\d{1,2}[/\-]\d{1,2}[/\-]\d{4})',
        ]:
            m = re.search(pat, texte, re.IGNORECASE | re.MULTILINE)
            if m:
                raw = re.sub(r'(?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)\s*', '', m.group(1), flags=re.IGNORECASE).strip()
                d = normaliser_date(raw)
                if d:
                    f['date_rando'] = d
                    break

    # ── NOM RANDO ──
    # Pattern explicite
    f['nom_rando'] = chercher(texte, [
        r'^\(in[t.]+ul[eé]\)\s*[\.…]*\s*(.+?)(?:\s*[\.…]+\s*)?$',
        r'^intitul[eé]\s*:?\s*(.+)',
    ])
    # Fallback : 2ème ligne si 1ère est une date
    if not f['nom_rando']:
        lignes = [l.strip() for l in texte.split('\n') if l.strip()]
        skip = {'lieu','heure','départ','depart','trajet','parking','kilom',
                'frais','données','donnees','durée','duree','organisateur',
                'équipement','equipement','inscription','animateur','téléphone',
                'telephone','email','covoiturage','randonnée','rando','distance',
                'dénivelé','denivele','difficulté','difficulte','niveau'}
        for ligne in lignes:
            # Ignorer dates
            if re.match(r'^\d{1,2}\s+\w+\s+\d{4}', ligne): continue
            if re.match(r'^\d{4}-\d{2}-\d{2}', ligne): continue
            if re.match(r'^(?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche)', ligne, re.IGNORECASE):
                # Ligne du type "Dimanche 17 MARS 2024" → passer, peut contenir le nom
                if re.search(r'\d{4}', ligne): continue
            # Ignorer les lignes qui sont clairement des champs de formulaire
            first_word = ligne.split()[0].lower().rstrip(':') if ligne.split() else ''
            if first_word in skip: continue
            if re.search(r'(?:' + '|'.join(skip) + r')\s*:', ligne, re.IGNORECASE): continue
            if re.match(r'^\(Jour', ligne, re.IGNORECASE): continue
            if 3 < len(ligne) < 120 and not re.match(r'^\d+\s*[xX€km]', ligne):
                f['nom_rando'] = ligne
                break
    # Dernier recours : nom du fichier
    if not f['nom_rando']:
        stem = Path(nom_fichier).stem
        stem = re.sub(r'^\d{4}[-_]\d{2}[-_]\d{2}[-_]?', '', stem)
        stem = re.sub(r'^\d{1,2}[-_]\w+[-_]\d{4}[-_]?', '', stem)
        f['nom_rando'] = stem.replace('_',' ').strip() or Path(nom_fichier).stem

    # ── ANIMATEUR ──
    f['animateur'] = chercher(texte, [
        r'[Oo]rganisateur(?:\s+de\s+la\s+randonn[ée]e?)?\s*:?\s*(.+)',
        r'[Aa]nimateur(?:-trice)?\s*:?\s*(.+)',
        r'[Ii]nscriptions?\s+aupr[eè]s\s+de\s+(.+?)(?:\s+par|\s+le|\s+avant|$)',
    ])
    if f['animateur']:
        # Retirer téléphone/email collé
        f['animateur'] = re.split(r'(?:t[ée]l[ée]phone|e-?mail|@|\d{2}[\s\.\-]\d{2}|\bpar\b|\bavant\b)', f['animateur'], flags=re.IGNORECASE)[0].strip()
        if f['animateur']: f['animateur'] = f['animateur'].rstrip(',:.').strip()
        # Rejeter si le résultat ressemble à une phrase et non un nom
        if re.search(r'(?:r[eé]serve|droit|modifier|circuit|itin)', f['animateur'], re.IGNORECASE):
            f['animateur'] = None
        f['animateur'] = f['animateur'].rstrip(',:.')

    # ── PARKING COVOIT ──
    f['parking_covoit'] = chercher(texte, [
        r'[Ll]ieu de rendez-vous\s*:?\s*(.+)',
        r'[Ll]ieu de rassemblement\s*:?\s*(.+)',
        r'[Pp]arking\s+(?:Henri\s+Dunant|[Hh]enri\s+[Dd]unant|[Dd]unant)',
        r'[Pp]arking\s+(?:RDV|rendez-vous)\s*:?\s*(.+)',
        r'(?:PARKING\s+D[ÉE]PART|PARKING\s+COVOIT)\s*:?\s*(.+)',
        r'[Rr](?:endez-vous|DV)\s*:?\s*(.+?)(?:\s+\d+h|\n|$)',
    ])
    # Cas particulier "Parking Henri Dunant" dans le texte
    if not f['parking_covoit'] and re.search(r'(?:henri\s+dunant|dunand|maison\s+des\s+associations)', texte, re.IGNORECASE):
        m = re.search(r'(parking\s+(?:henri\s+dunant|dunand|de\s+la\s+maison\s+des\s+associations)[^\n]*)', texte, re.IGNORECASE)
        if m: f['parking_covoit'] = m.group(1).strip()
    if f['parking_covoit']:
        f['parking_covoit'] = f['parking_covoit'][:200]

    # ── HEURE RV ──
    heure_raw = chercher(texte, [
        r'[Hh]eure\s+(?:du\s+)?(?:rendez-vous|RV?|rassemblement)[^:]*:?\s*(\d{1,2}h\d{0,2})',
        r'[Rr](?:assemblement|DV|endez-vous)\s*:?\s*(\d{1,2}h\d{0,2})',
        r'[Hh]eure\s+RV?\s*:?\s*(\d{1,2}h\d{0,2})',
        r'RV\s*:\s*(\d{1,2}h\d{0,2})',
    ])
    f['heure_rv'] = normaliser_heure(heure_raw)

    # ── PARKING DÉPART ──
    f['parking_depart'] = chercher(texte, [
        r'[Ll]ieu de parking et d[ée]part[^:]*:?\s*(.+)',
        r'[Pp]arking\s+(?:et\s+)?d[ée]part\s*:?\s*(.+)',
        r'(?:PARKING\s+D[ÉE]PART|Parking\s+:)\s*(.+)',
        r'[Ss]e\s+garer?\s+(.+)',
    ])
    if f['parking_depart']:
        f['parking_depart'] = f['parking_depart'][:200]

    # ── GPS ──
    gps = chercher(texte, [
        r'GPS\s+(?:Maps?\s*:?\s*)(.+)',
        r'(?:Lat|lat)\s*:?\s*([\d,\.]+)[^\d].*?(?:Long?|lon)\s*:?\s*([\d,\.]+)',
        r'(\d{2,3}[.,]\d+\s*[NS]?\s*/?\s*\d{1,3}[.,]\d+\s*[EW]?)',
        r'GPS\s*:?\s*([0-9.,\s/]+)',
    ])
    # GPS sous forme "43.7523228 N / 4.793681 E" dans le texte
    if not gps:
        m = re.search(r'([\d,\.]+)\s*N\s*/?\s*([\d,\.]+)\s*E', texte)
        if m: gps = f"{m.group(1)},{m.group(2)}"
    # GPS sous forme "Lat : 43,7470/ Long : 4,9214"
    if not gps:
        m = re.search(r'[Ll]at\s*:?\s*([\d,\.]+).*?[Ll]on[g]?\s*:?\s*([\d,\.]+)', texte, re.DOTALL)
        if m: gps = f"{m.group(1).replace(',','.')},{m.group(2).replace(',','.')}"
    f['gps'] = gps

    # ── DISTANCE ──
    f['distance'] = chercher(texte, [
        r'[Dd]istance\s*:?\s*[≃~]?\s*([\d,\.]+)\s*km',
        r'[Ll]ongueur\s*:?\s*[+/\-]*\s*([\d,\.]+)\s*km',
        r'(?:^|\n)\s*([\d,\.]+)\s*km\b',
    ], float)

    # ── DÉNIVELÉ ──
    f['denivele'] = chercher(texte, [
        r'[Dd][ée]nivel[ée]\s*(?:positif|[+\+])?\s*:?\s*[+/\-≃~]*\s*([\d\s\u202f]+)\s*m',
        r'[Dd][ée]nivel[ée]\s*:?\s*[+/\-≃~]*\s*([\d\s\u202f]+)\s*m',
    ], int)
    if f['denivele']:
        try: f['denivele'] = int(str(f['denivele']).replace(' ','').replace('\u202f',''))
        except: pass

    # ── DURÉE ──
    f['duree'] = chercher(texte, [
        r'[Dd]ur[ée]e\s*:?\s*[≃~+/\-]*\s*([\d,\.]+\s*h(?:\s*\d{0,2})?)',
        r'[Dd]ur[ée]e\s*:?\s*[≃~+/\-]*\s*([\d,\.]+)\s*heures?',
        r'[Dd]ur[ée]e\s*:?\s*([\dhHmM\s,\.]+)',
    ])
    if f['duree']:
        f['duree'] = f['duree'].strip().rstrip('.')

    # ── IBP ──
    f['ibp'] = chercher(texte, [
        r'[Ii](?:ndice\s+)?[Bb][Pp]\s*[:\.]?\s*([\d,\.]+)',
        r'IBP\s+([\d,\.]+)',
        r'IPB\s+([\d,\.]+)',  # typo courante
    ], int)

    # ── EFFORT ──
    f['effort'] = chercher(texte, [
        r'[Ee]ffort\s+(\d)\s*/?\s*5',
        r'[Ee]ffort\s*[:\.]?\s*(\d)',
    ], int)

    # ── TECHNICITÉ ──
    f['technicite'] = chercher(texte, [
        r'[Tt]echnici(?:t[ée]|te)\s+(\d)\s*/?\s*5',
        r'[Tt]echnici(?:t[ée]|te)\s*[:\.]?\s*(\d)',
    ], int)

    # ── RISQUE ──
    f['risque'] = chercher(texte, [
        r'[Rr]isque\s+(\d)\s*/?\s*5',
        r'[Rr]isque\s*[:\.]?\s*(\d)',
    ], int)

    # ── COÛTS ──
    couts = chercher(texte, [
        r'[Ff]rais de covoiturage[^:]*:?\s*(.{5,100})',
        r'[Cc]ovoiturage\s*=\s*(.{5,80})',
        r'\d+\s*(?:km)?\s*[xX]\s*0[,\.]\d+\s*€?\s*=\s*([\d,\.\s€]+(?:/voiture)?)',
    ])
    if couts:
        f['couts'] = couts.strip()[:200]

    # ── REMARQUES ──
    rem = chercher(texte, [
        r'[Rr]emarques?\s+sur\s+la\s+rando(?:nn[ée]e?)?\s*:?\s*(.+?)(?=\n[A-ZL\'Oo]|\n\n|$)',
        r'[Rr]emarques?\s*:?\s*(.+?)(?=\n[A-ZL\'Oo]|\n\n|$)',
        r'[Dd]ifficulté(?:s)?\s+(?:et\s+particularités?)?\s*:?\s*(.+?)(?=\n[A-Z]|\n\n|$)',
    ])
    if rem:
        # Nettoyer et limiter
        rem = re.sub(r'\s+', ' ', rem).strip()
        f['remarques'] = rem[:500]

    return f

# ══ SUPABASE ══════════════════════════════════════════════════

def inserer(fiche):
    r = requests.post(
        f"{SUPABASE_URL}/rest/v1/fiches",
        headers={
            "Content-Type": "application/json",
            "apikey": SUPABASE_KEY,
            "Authorization": f"Bearer {SUPABASE_KEY}",
            "Prefer": "return=minimal"
        },
        json=fiche,
        timeout=10
    )
    return r.status_code in (200, 201), r.text

# ══ MAIN ══════════════════════════════════════════════════════

def main():
    args = [a for a in sys.argv[1:] if not a.startswith('--')]
    if not args:
        print("Usage: python3 import_fdr.py /dossier [--prod]")
        sys.exit(1)

    dossier = Path(args[0])
    exts = {'.odt', '.docx', '.doc', '.pdf'}
    fichiers = sorted([f for f in dossier.iterdir() if f.suffix.lower() in exts])

    print(f"\n🥾 Import FDR → Supabase")
    print(f"📁 {dossier}  ({len(fichiers)} fichiers)")
    print(f"🔧 Mode : {'PRODUCTION' if MODE_PROD else 'TEST (ajoutez --prod pour insérer)'}\n")

    ok = ko = skip = 0
    resultats = []

    for chemin in fichiers:
        nom = chemin.name
        ext = chemin.suffix.lower()
        print(f"{'─'*55}")
        print(f"📄 {nom}")

        # Extraction
        if ext == '.odt':    texte = extraire_odt(chemin)
        elif ext == '.docx': texte = extraire_docx(chemin)
        elif ext == '.doc':  texte = extraire_doc(chemin)
        elif ext == '.pdf':  texte = extraire_pdf(chemin)
        else: texte = ""

        if not texte or len(texte.strip()) < 30:
            print(f"  ⏭  Texte vide (PDF image ou fichier vide)")
            skip += 1
            continue

        fiche = parser(texte, nom)
        resultats.append({'fichier': nom, 'fiche': fiche})

        print(f"  📅 {fiche['date_rando'] or '—'}")
        print(f"  🏔  {fiche['nom_rando'] or '—'}")
        print(f"  👤 {fiche['animateur'] or '—'}")
        print(f"  🅿️  RDV: {fiche['parking_covoit'] or '—'}  ⏰ {fiche['heure_rv'] or '—'}")
        print(f"  📏 {fiche['distance'] or '—'} km  ⛰  {fiche['denivele'] or '—'} m  ⏱  {fiche['duree'] or '—'}")
        print(f"  📊 IBP:{fiche['ibp'] or '—'}  Effort:{fiche['effort'] or '—'}  Tech:{fiche['technicite'] or '—'}  Risque:{fiche['risque'] or '—'}")

        if MODE_PROD:
            success, msg = inserer(fiche)
            if success:
                print(f"  ✅ Inséré")
                ok += 1
            else:
                print(f"  ❌ Erreur: {msg[:80]}")
                ko += 1
        else:
            print(f"  ✅ [TEST]")
            ok += 1

    print(f"\n{'═'*55}")
    print(f"✅ {ok}  ❌ {ko}  ⏭  {skip}")
    if not MODE_PROD:
        print(f"\n💡 Pour insérer : python3 import_fdr.py {args[0]} --prod")

    # Sauvegarder le rapport JSON
    with open('/home/claude/rapport_import.json', 'w', encoding='utf-8') as jf:
        json.dump(resultats, jf, ensure_ascii=False, indent=2, default=str)
    print(f"📋 Rapport JSON : /home/claude/rapport_import.json")

if __name__ == "__main__":
    main()
